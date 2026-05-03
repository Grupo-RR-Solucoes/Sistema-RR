import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"C:\Users\diego\Documents\Codex\2026-04-20-files-mentioned-by-the-user-sistema\repo\Sistema-RR-main")
OUTPUT_DIR = ROOT / "runtime" / "outputs"
JSON_PATH = OUTPUT_DIR / "auditoria-historica-avista-prt-2022-12-2026-03.json"
XLSX_PATH = OUTPUT_DIR / "auditoria-historica-avista-prt-2022-12-2026-03.xlsx"
DOCX_PATH = OUTPUT_DIR / "minuta-questionamento-divergencias-avista-prt.docx"
MD_PATH = OUTPUT_DIR / "minuta-questionamento-divergencias-avista-prt.md"


def brl(value: float) -> str:
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def period_label(period: str) -> str:
    year, month = period.split("-")
    months = {
        "01": "jan",
        "02": "fev",
        "03": "mar",
        "04": "abr",
        "05": "mai",
        "06": "jun",
        "07": "jul",
        "08": "ago",
        "09": "set",
        "10": "out",
        "11": "nov",
        "12": "dez",
    }
    return f"{months.get(month, month)}/{year[2:]}"


with JSON_PATH.open("r", encoding="utf-8") as f:
    payload = json.load(f)

meta = payload["meta"]
cover_rows = payload["coverRows"]
monthly_rows = payload["monthlyRows"]
divergent_rows = payload["divergentRows"]
pending_rows = payload["pendingRows"]

top_months = sorted(
    monthly_rows,
    key=lambda row: float(row["excessoAvista"] or 0)
    + float(row["faltaDiferido"] or 0)
    + float(row["faltaPrt"] or 0),
    reverse=True,
)[:15]

company_blocks = []
for row in cover_rows:
    monthly = [
        item
        for item in monthly_rows
        if item["empresaCnpj"] == row["empresaCnpj"]
        and int(item["operacoesDivergentes"] or 0) > 0
    ]
    avista_total = round(sum(float(item["excessoAvista"] or 0) for item in monthly), 2)
    diferido_total = round(sum(float(item["faltaDiferido"] or 0) for item in monthly), 2)
    prt_total = round(sum(float(item["faltaPrt"] or 0) for item in monthly), 2)
    top_company_months = sorted(
        monthly,
        key=lambda item: float(item["excessoAvista"] or 0)
        + float(item["faltaDiferido"] or 0)
        + float(item["faltaPrt"] or 0),
        reverse=True,
    )[:5]

    company_blocks.append(
        {
            "name": row["empresaNome"],
            "cnpj": row["empresaCnpj"],
            "firstPeriod": row["primeiraCompetencia"],
            "lastPeriod": row["ultimaCompetencia"],
            "operationsCash": row["operacoesCash"],
            "operationsAudited": row["operacoesAuditoriaCompleta"],
            "divergences": row["divergencias"],
            "avistaTotal": avista_total,
            "diferidoTotal": diferido_total,
            "prtTotal": prt_total,
            "topMonths": top_company_months,
        }
    )

company_blocks.sort(
    key=lambda item: item["avistaTotal"] + item["diferidoTotal"] + item["prtTotal"],
    reverse=True,
)

status_counts = defaultdict(int)
for row in divergent_rows:
    status_counts[row["statusGeral"]] += 1

document = Document()
styles = document.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MINUTA DE QUESTIONAMENTO\nDIVERGENCIAS DE A VISTA, DIFERIDO E PRT")
run.bold = True
run.font.size = Pt(15)

sub = document.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Documento de apoio para solicitacao de memoria de calculo, validacao dos percentuais a vista e regularizacao das divergencias financeiras"
).italic = True

document.add_paragraph("")

p = document.add_paragraph()
p.add_run("Objeto. ").bold = True
p.add_run(
    "Solicitar a conferencia formal das divergencias identificadas na auditoria historica de a vista, diferido total e PRT mensal, "
    "com base nos fechamentos mensais importados para todas as empresas no intervalo de dezembro/2022 a marco/2026."
)

p = document.add_paragraph()
p.add_run("Criterio da auditoria. ").bold = True
p.add_run(
    "Cada linha de A Vista foi confrontada com a OPP/TRP vigente da competencia. "
    "Quando a regra do mes admitia mais de uma faixa valida de a vista e a base nao trazia a faixa/meta explicitamente, "
    "o enquadramento foi tratado como intervalo valido. "
    "O diferido total foi tratado como a comissao total da operacao menos o a vista. "
    "O PRT mensal foi tratado como o diferido total parcelado igualmente pelo numero total de parcelas do contrato. "
    "As linhas sem PRT suficiente para reconstruir a carteira foram mantidas como pendencia, e nao como divergencia fechada."
)

p = document.add_paragraph()
p.add_run("Escopo excluido. ").bold = True
p.add_run(
    "Arquivos e operacoes de Consorcio e BrasilCAP foram desconsiderados desta auditoria."
)

p = document.add_paragraph()
p.add_run("Base de apoio. ").bold = True
p.add_run(f"Planilha analitica consolidada: {XLSX_PATH}")

p = document.add_paragraph()
p.add_run("Observacao de base. ").bold = True
p.add_run(
    "O fechamento de 12/2023 da empresa RR SOLUCOES LTDA (48...) permanece zerado no arquivo de origem atualmente localizado, "
    "de modo que esta competencia segue dependente de eventual arquivo substituto para fechamento definitivo."
)

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("1. Resumo executivo da auditoria").bold = True

summary_bullets = [
    f"Operacoes A Vista analisadas: {meta['operationsCash']}",
    f"Operacoes classificadas como divergentes: {meta['operationsDivergent']}",
    f"Operacoes mantidas como pendencia por falta de PRT suficiente para reconstruir a carteira: {meta['operationsPending']}",
    f"Empresas auditadas: {len(cover_rows)}",
]

for item in summary_bullets:
    p = document.add_paragraph()
    p.add_run("• ").bold = True
    p.add_run(item)

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("2. Competencias mais criticas").bold = True

table = document.add_table(rows=1, cols=8)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Empresa"
hdr[1].text = "CNPJ"
hdr[2].text = "Competencia"
hdr[3].text = "Politica a vista"
hdr[4].text = "Operacoes divergentes"
hdr[5].text = "Excesso de a vista"
hdr[6].text = "Falta de diferido"
hdr[7].text = "Falta de PRT"

for row in top_months:
    cells = table.add_row().cells
    cells[0].text = row["empresaNome"]
    cells[1].text = row["empresaCnpj"]
    cells[2].text = period_label(row["competencia"])
    cells[3].text = row["politicaAvistaMes"]
    cells[4].text = str(row["operacoesDivergentes"])
    cells[5].text = brl(float(row["excessoAvista"] or 0))
    cells[6].text = brl(float(row["faltaDiferido"] or 0))
    cells[7].text = brl(float(row["faltaPrt"] or 0))

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("3. Consolidado por empresa").bold = True

for company in company_blocks:
    p = document.add_paragraph()
    p.add_run(f"{company['name']} ({company['cnpj']})").bold = True
    p.add_run(
        f": periodo auditado de {period_label(company['firstPeriod'])} a {period_label(company['lastPeriod'])}, "
        f"{company['operationsCash']} operacoes de A Vista, "
        f"{company['operationsAudited']} operacoes com auditoria completa de diferido/PRT e "
        f"{company['divergences']} operacoes divergentes."
    )

    p = document.add_paragraph()
    p.add_run("• ").bold = True
    p.add_run(
        f"Excesso acumulado de a vista identificado: {brl(company['avistaTotal'])}; "
        f"falta acumulada de diferido: {brl(company['diferidoTotal'])}; "
        f"falta acumulada de PRT mensal: {brl(company['prtTotal'])}."
    )

    for month in company["topMonths"]:
        bullet = document.add_paragraph()
        bullet.add_run("• ").bold = True
        bullet.add_run(
            f"{period_label(month['competencia'])}: politica do mes {month['politicaAvistaMes']}, "
            f"{month['operacoesDivergentes']} operacoes divergentes, "
            f"excesso de a vista {brl(float(month['excessoAvista'] or 0))}, "
            f"falta de diferido {brl(float(month['faltaDiferido'] or 0))} e "
            f"falta de PRT {brl(float(month['faltaPrt'] or 0))}."
        )

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("4. Questionamentos objetivos a serem respondidos").bold = True

questions = [
    "Solicitamos a memoria de calculo da comissao a vista por operacao nas competencias questionadas, com indicacao da OPP/TRP aplicada, faixa de producao/meta considerada e percentual a vista efetivamente utilizado em cada operacao.",
    "Solicitamos a memoria de calculo da comissao total de credito por operacao, demonstrando o percentual total utilizado e o valor de diferido total que deveria permanecer em carteira apos o a vista.",
    "Solicitamos o espelho da carteira PRT por operacao e por competencia, contendo no minimo: numero da operacao, data final, quantidade de parcelas total, quantidade de parcelas pagas, valor unitario da parcela e situacao atual da operacao.",
    "Solicitamos a relacao formal das operacoes baixadas por liquidacao, cancelamento, renovacao ou qualquer outro motivo que tenha interrompido o PRT, com a respectiva data de baixa e o numero da operacao.",
    "Solicitamos confirmar, por competencia, qual tabela OPP/TRP estava vigente e qual faixa/meta foi aplicada para enquadramento do a vista em cada empresa.",
    "Solicitamos esclarecer as competencias em que o a vista recebido se mostrou acima da faixa valida, com consequente reducao do diferido total e do PRT mensal.",
    "Solicitamos encaminhar eventual fechamento valido substitutivo da competencia 12/2023 da empresa RR SOLUCOES LTDA (48...), caso exista versao correta diferente do arquivo atualmente localizado.",
]

for item in questions:
    p = document.add_paragraph()
    p.add_run("• ").bold = True
    p.add_run(item)

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("5. Solicitacao de regularizacao e retorno").bold = True

p = document.add_paragraph()
p.add_run(
    "Diante dos apontamentos acima, solicitamos a conferencia formal das competencias indicadas, o encaminhamento da memoria de calculo por operacao e, "
    "se confirmadas as divergencias, a regularizacao financeira dos valores de diferido e PRT pagos a menor."
)

p = document.add_paragraph()
p.add_run(
    "Permanecemos a disposicao para confrontar operacao por operacao com base na planilha anexa de auditoria historica."
)

document.add_paragraph("")
p = document.add_paragraph()
p.add_run("Anexo de apoio: ").bold = True
p.add_run(str(XLSX_PATH))

document.save(str(DOCX_PATH))

md_lines = [
    "# Minuta de Questionamento - Divergencias de A Vista, Diferido e PRT",
    "",
    "## Objeto",
    "Solicitar a conferencia formal das divergencias identificadas na auditoria historica de a vista, diferido total e PRT mensal, com base nos fechamentos mensais importados para todas as empresas no intervalo de dezembro/2022 a marco/2026.",
    "",
    "## Criterio da auditoria",
    "Cada linha de A Vista foi confrontada com a OPP/TRP vigente da competencia. Quando a regra do mes admitia mais de uma faixa valida de a vista e a base nao trazia a faixa/meta explicitamente, o enquadramento foi tratado como intervalo valido. O diferido total foi tratado como a comissao total da operacao menos o a vista. O PRT mensal foi tratado como o diferido total parcelado igualmente pelo numero total de parcelas do contrato. As linhas sem PRT suficiente para reconstruir a carteira foram mantidas como pendencia, e nao como divergencia fechada.",
    "",
    "## Escopo excluido",
    "Consorcio e BrasilCAP nao foram considerados nesta auditoria.",
    "",
    "## Base de apoio",
    f"Planilha analitica consolidada: `{XLSX_PATH}`",
    "",
    "## Observacao de base",
    "O fechamento de 12/2023 da empresa RR SOLUCOES LTDA (48...) permanece zerado no arquivo de origem atualmente localizado, de modo que esta competencia segue dependente de eventual arquivo substituto para fechamento definitivo.",
    "",
    "## Resumo executivo da auditoria",
    f"- Operacoes A Vista analisadas: {meta['operationsCash']}",
    f"- Operacoes classificadas como divergentes: {meta['operationsDivergent']}",
    f"- Operacoes mantidas como pendencia por falta de PRT suficiente para reconstruir a carteira: {meta['operationsPending']}",
    f"- Empresas auditadas: {len(cover_rows)}",
    "",
    "## Competencias mais criticas",
    "",
    "| Empresa | CNPJ | Competencia | Politica a vista | Operacoes divergentes | Excesso de a vista | Falta de diferido | Falta de PRT |",
    "|---|---|---|---|---:|---:|---:|---:|",
]

for row in top_months:
    md_lines.append(
        f"| {row['empresaNome']} | {row['empresaCnpj']} | {period_label(row['competencia'])} | {row['politicaAvistaMes']} | {row['operacoesDivergentes']} | {brl(float(row['excessoAvista'] or 0))} | {brl(float(row['faltaDiferido'] or 0))} | {brl(float(row['faltaPrt'] or 0))} |"
    )

md_lines.extend(["", "## Consolidado por empresa", ""])

for company in company_blocks:
    md_lines.append(
        f"- **{company['name']} ({company['cnpj']})**: periodo auditado de {period_label(company['firstPeriod'])} a {period_label(company['lastPeriod'])}, {company['operationsCash']} operacoes de A Vista, {company['operationsAudited']} operacoes com auditoria completa de diferido/PRT e {company['divergences']} operacoes divergentes."
    )
    md_lines.append(
        f"  - Excesso acumulado de a vista identificado: {brl(company['avistaTotal'])}; falta acumulada de diferido: {brl(company['diferidoTotal'])}; falta acumulada de PRT mensal: {brl(company['prtTotal'])}."
    )
    for month in company["topMonths"]:
        md_lines.append(
            f"  - {period_label(month['competencia'])}: politica do mes {month['politicaAvistaMes']}, {month['operacoesDivergentes']} operacoes divergentes, excesso de a vista {brl(float(month['excessoAvista'] or 0))}, falta de diferido {brl(float(month['faltaDiferido'] or 0))} e falta de PRT {brl(float(month['faltaPrt'] or 0))}."
        )

md_lines.extend(["", "## Questionamentos objetivos a serem respondidos", ""])
for item in questions:
    md_lines.append(f"- {item}")

md_lines.extend(
    [
        "",
        "## Solicitacao de regularizacao e retorno",
        "",
        "Diante dos apontamentos acima, solicitamos a conferencia formal das competencias indicadas, o encaminhamento da memoria de calculo por operacao e, se confirmadas as divergencias, a regularizacao financeira dos valores de diferido e PRT pagos a menor.",
        "",
        "Permanecemos a disposicao para confrontar operacao por operacao com base na planilha anexa de auditoria historica.",
        "",
        f"**Anexo de apoio:** `{XLSX_PATH}`",
    ]
)

MD_PATH.write_text("\n".join(md_lines), encoding="utf-8")

print(
    json.dumps(
        {
            "docx": str(DOCX_PATH),
            "md": str(MD_PATH),
            "xlsx": str(XLSX_PATH),
            "companies": len(company_blocks),
            "topMonths": len(top_months),
            "divergentRows": len(divergent_rows),
            "pendingRows": len(pending_rows),
        },
        ensure_ascii=False,
        indent=2,
    )
)
