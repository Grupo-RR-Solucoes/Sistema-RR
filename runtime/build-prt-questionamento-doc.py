import json
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(r"C:\Users\diego\Documents\Codex\2026-04-20-files-mentioned-by-the-user-sistema\repo\Sistema-RR-main")
OUTPUT_DIR = ROOT / "runtime" / "outputs"
JSON_PATH = OUTPUT_DIR / "prt-divergencias-detalhado.json"
DOCX_PATH = OUTPUT_DIR / "minuta-questionamento-divergencias-prt.docx"
MD_PATH = OUTPUT_DIR / "minuta-questionamento-divergencias-prt.md"
XLSX_PATH = OUTPUT_DIR / "prt-divergencias-detalhado.xlsx"


def brl(value: float) -> str:
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def period_label(period: str) -> str:
    year, month = period.split("-")
    months = {
        "01": "jan",
        "02": "fev",
        "03": "mar",
        "04": "abr",
        "05": "mai",
        "06": "jun",
        "07": "jul",
        "08": "ago",
        "09": "set",
        "10": "out",
        "11": "nov",
        "12": "dez",
    }
    return f"{months.get(month, month)}/{year[2:]}"


with JSON_PATH.open("r", encoding="utf-8") as f:
    payload = json.load(f)

summary_rows = payload["summaryRows"]
detail_rows = payload["detailRows"]

top_rows = sorted(summary_rows, key=lambda row: row["diferenca"], reverse=True)[:12]

by_company = defaultdict(lambda: {"total": 0.0, "months": 0, "top": []})
for row in summary_rows:
    key = (row["empresaNome"], row["empresaCnpj"])
    by_company[key]["total"] += float(row["diferenca"] or 0)
    by_company[key]["months"] += 1
    by_company[key]["top"].append(row)

company_blocks = []
for (name, cnpj), info in by_company.items():
    months = sorted(info["top"], key=lambda row: row["diferenca"], reverse=True)[:5]
    company_blocks.append(
        {
            "name": name,
            "cnpj": cnpj,
            "total": round(info["total"], 2),
            "months": info["months"],
            "top": months,
        }
    )

company_blocks.sort(key=lambda item: item["total"], reverse=True)

document = Document()

styles = document.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MINUTA DE QUESTIONAMENTO\nDIVERGENCIAS DE PRT / DIFERIDO")
run.bold = True
run.font.size = Pt(15)

sub = document.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Documento de apoio para solicitacao de memoria de calculo, validacao e regularizacao de valores").italic = True

document.add_paragraph("")

p = document.add_paragraph()
p.add_run("Objeto. ").bold = True
p.add_run(
    "Solicitar a conferencia formal das divergencias identificadas no recebimento de PRT/diferido, "
    "com base na auditoria interna realizada a partir dos fechamentos mensais historicos ja importados."
)

p = document.add_paragraph()
p.add_run("Criterio utilizado na auditoria. ").bold = True
p.add_run(
    "Cada linha de PRT foi tratada como parcela mensal recorrente em valor igual, seguindo o numero total de parcelas do contrato e o numero de parcelas ja pagas. "
    "O valor esperado em cada competencia corresponde a soma das parcelas que deveriam estar em curso naquele mes, "
    "sendo interrompidas apenas quando identificada ocorrencia de liquidacao, cancelamento ou renovacao."
)

p = document.add_paragraph()
p.add_run("Base de apoio. ").bold = True
p.add_run(
    f"A relacao analitica por operacao foi consolidada na planilha anexa: {XLSX_PATH}"
)

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("1. Principais competencias com divergencia identificada").bold = True

table = document.add_table(rows=1, cols=6)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Empresa"
hdr[1].text = "CNPJ"
hdr[2].text = "Competencia"
hdr[3].text = "PRT recebido no mes"
hdr[4].text = "Valor divergente identificado"
hdr[5].text = "Operacoes divergentes"

for row in top_rows:
    cells = table.add_row().cells
    cells[0].text = row["empresaNome"]
    cells[1].text = row["empresaCnpj"]
    cells[2].text = period_label(row["competencia"])
    cells[3].text = brl(float(row["actualPrt"] or 0))
    cells[4].text = brl(float(row["diferenca"] or 0))
    cells[5].text = str(row["operacoesDivergentes"])

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("2. Consolidado por empresa").bold = True

for company in company_blocks:
    p = document.add_paragraph(style=None)
    p.add_run(f"{company['name']} ({company['cnpj']})").bold = True
    p.add_run(
        f": divergencia acumulada identificada de {brl(company['total'])}, distribuida em {company['months']} competencias com apontamento."
    )
    for month in company["top"]:
        bullet = document.add_paragraph(style=None)
        bullet.style = document.styles["Normal"]
        bullet.add_run("• ").bold = True
        bullet.add_run(
            f"{period_label(month['competencia'])}: recebido {brl(float(month['actualPrt'] or 0))}, "
            f"divergencia identificada {brl(float(month['diferenca'] or 0))}, "
            f"{month['operacoesDivergentes']} operacoes."
        )

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("3. Questionamentos objetivos a serem respondidos").bold = True

questions = [
    "Solicitamos a memoria de calculo do PRT pago em cada competencia listada, discriminada por operacao, com indicacao do valor da parcela, numero da parcela paga e saldo remanescente.",
    "Solicitamos o espelho da carteira de PRT ativa por competencia, contendo no minimo: numero da operacao, data final, quantidade de parcelas total, quantidade de parcelas pagas, valor unitario da parcela e situacao atual da operacao.",
    "Solicitamos a justificativa formal para as competencias em que o valor recebido ficou abaixo do valor identificado em nossa auditoria interna.",
    "Solicitamos a relacao das operacoes baixadas por liquidacao, cancelamento, renovacao ou qualquer outro motivo que tenha interrompido o PRT, com a respectiva data de baixa e o numero da operacao.",
    "Solicitamos informar se houve alteracao de regra comercial, OPP ou TRP aplicavel aos periodos questionados e, em caso positivo, encaminhar a tabela vigente em cada competencia.",
    "Solicitamos confirmar se existe algum arquivo complementar, ajuste manual ou rotina de compensacao que nao esteja refletida nos fechamentos mensais ja enviados.",
]

for item in questions:
    p = document.add_paragraph()
    p.add_run("• ").bold = True
    p.add_run(item)

document.add_paragraph("")
heading = document.add_paragraph()
heading.add_run("4. Solicitação de regularização e retorno").bold = True

p = document.add_paragraph()
p.add_run(
    "Diante das divergencias acima, solicitamos a conferencia formal das competencias apontadas, "
    "o encaminhamento da memoria de calculo por operacao e, se confirmado o pagamento a menor, "
    "a respectiva regularizacao financeira dos valores devidos."
)

p = document.add_paragraph()
p.add_run(
    "Permanecemos a disposicao para confrontar as operacoes individualmente, tendo como base a planilha detalhada anexa."
)

document.add_paragraph("")
p = document.add_paragraph()
p.add_run("Anexo de apoio: ").bold = True
p.add_run(str(XLSX_PATH))

document.save(str(DOCX_PATH))

md_lines = [
    "# Minuta de Questionamento - Divergências de PRT / Diferido",
    "",
    "## Objeto",
    "Solicitar a conferência formal das divergências identificadas no recebimento de PRT/diferido, com base na auditoria interna realizada a partir dos fechamentos mensais históricos já importados.",
    "",
    "## Critério utilizado na auditoria",
    "Cada linha de PRT foi tratada como parcela mensal recorrente em valor igual, seguindo o número total de parcelas do contrato e o número de parcelas já pagas. O valor esperado em cada competência corresponde à soma das parcelas que deveriam estar em curso naquele mês, sendo interrompidas apenas quando identificada ocorrência de liquidação, cancelamento ou renovação.",
    "",
    "## Base de apoio",
    f"Planilha analítica por operação: `{XLSX_PATH}`",
    "",
    "## Principais competências com divergência identificada",
    "",
    "| Empresa | CNPJ | Competência | PRT recebido no mês | Valor divergente identificado | Operações divergentes |",
    "|---|---|---:|---:|---:|---:|",
]

for row in top_rows:
    md_lines.append(
        f"| {row['empresaNome']} | {row['empresaCnpj']} | {period_label(row['competencia'])} | {brl(float(row['actualPrt'] or 0))} | {brl(float(row['diferenca'] or 0))} | {row['operacoesDivergentes']} |"
    )

md_lines.extend(
    [
        "",
        "## Consolidado por empresa",
        "",
    ]
)

for company in company_blocks:
    md_lines.append(
        f"- **{company['name']} ({company['cnpj']})**: divergência acumulada identificada de {brl(company['total'])}, distribuída em {company['months']} competências com apontamento."
    )
    for month in company["top"]:
        md_lines.append(
            f"  - {period_label(month['competencia'])}: recebido {brl(float(month['actualPrt'] or 0))}, divergência identificada {brl(float(month['diferenca'] or 0))}, {month['operacoesDivergentes']} operações."
        )

md_lines.extend(
    [
        "",
        "## Questionamentos objetivos a serem respondidos",
        "",
    ]
)

for item in questions:
    md_lines.append(f"- {item}")

md_lines.extend(
    [
        "",
        "## Solicitação de regularização e retorno",
        "",
        "Diante das divergências acima, solicitamos a conferência formal das competências apontadas, o encaminhamento da memória de cálculo por operação e, se confirmado o pagamento a menor, a respectiva regularização financeira dos valores devidos.",
        "",
        "Permanecemos à disposição para confrontar as operações individualmente, tendo como base a planilha detalhada anexa.",
        "",
        f"**Anexo de apoio:** `{XLSX_PATH}`",
    ]
)

MD_PATH.write_text("\n".join(md_lines), encoding="utf-8")

print(
    json.dumps(
        {
            "docx": str(DOCX_PATH),
            "md": str(MD_PATH),
            "xlsx": str(XLSX_PATH),
            "topRows": len(top_rows),
            "companies": len(company_blocks),
        },
        ensure_ascii=False,
        indent=2,
    )
)
